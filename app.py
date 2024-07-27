import json
import re
import PyPDF2
import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from docx import Document#exceptions
from io import BytesIO
import fitz  
st.set_page_config(page_title="Seo copywriter", layout="wide", initial_sidebar_state="expanded")
from docx.shared import Pt
import os
st.markdown('<h1>SEO CONTENT WRITER</h1>', unsafe_allow_html=True)
from langchain_core.messages import HumanMessage
import anthropic
from dotenv import load_dotenv


file_path = os.path.join("prompt", "prompt.txt")
extracting_prompt = os.path.join("prompt", "extraction_prompt.txt")
api_key= os.getenv("ANTHROPIC_API_KEY")

# Check if the file exists
if os.path.exists(file_path):
    with open(file_path, 'r') as prompt_file:
        prompt = prompt_file.read()

if os.path.exists(file_path):
    with open(file_path, 'r') as prompt_file:
        extraction_prompt = prompt_file.read()
    
def get_checkbox_states(pdf):
    checkbox_states = []
    pdf_reader = PyPDF2.PdfReader(pdf)
    fields = pdf_reader.get_fields()
    if fields:
        checkboxes = {k: v.get('/V', 'Off') for k, v in fields.items() if v.get('/FT') == '/Btn'}
        checkbox_states.append(checkboxes)
    return checkbox_states

def get_pdf_text(pdf):
    texts = []
    text = ""
    pdf_document = fitz.open(stream=pdf.read(), filetype="pdf")
    for page_num in range(0, pdf_document.page_count):  
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    # Remove placeholders (underscores)
    text = re.sub(r'_{2,}', '', text)
    #print(get_checkbox_states(pdf))
    return text
def get_pdf_text_cahier(pdf):
    texts = []
    text = ""
    pdf_document = fitz.open(stream=pdf.read(), filetype="pdf")
    for page_num in range(0, pdf_document.page_count):  
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    # Remove placeholders (underscores)
    text = re.sub(r'_{2,}', '', text)
    #print(get_checkbox_states(pdf))
    return text


def get_page_text(pdf):
    pdf_reader = PdfReader(pdf)
    pdf_text = ""
    for page in pdf_reader.pages:
        pdf_text += page.extract_text()
    return pdf_reader

def create_word_doc(texts, titles):
    doc = Document()
    
    # Add a header with big text
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run("SEO Content")
    run.font.size = Pt(24)
    run.bold = True
    
    for i, (text, title_text) in enumerate(zip(texts, titles)):
        # Add a bold title before each text
        title = doc.add_paragraph()
        run = title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph(text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def get_vectorstore(texts):
    if not texts:
        raise ValueError("No texts to process for vector store creation.")
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_texts(texts=texts, embedding=embeddings)
    return vectorstore
client = anthropic.Anthropic(
    api_key=api_key
)
def extract_data(user_question):
    
    system_prompt = extraction_prompt
    response = client.messages.create(
        model="claude-3-5-sonnet-20240620",
        max_tokens=4096,
        system=system_prompt,
        messages=[
            {"role": "user", "content": user_question}
        ]
    )
    answer=response.content[0].text
    return answer
def handle_userinput(user_question):
    
    system_prompt = prompt

# Create a message

    response = client.messages.create(
        model="claude-3-5-sonnet-20240620",
        max_tokens=4096,
        system=system_prompt,
        messages=[
            {"role": "user", "content": user_question}
        ]
    )
    answer=response.content[0].text
    return answer

def main():


    load_dotenv()

    st.subheader("Dynamic Structure")

    num_pages = st.number_input("Number of pages:", min_value=1, max_value=10, value=1, step=1)
    pdf_templates_folder = 'pdf_templates'
    options = [f for f in os.listdir(pdf_templates_folder) if f.endswith('.pdf')]
    page_details = []

    for i in range(num_pages):
        st.write(f"### Page {i+1}")
        text_input_1 = st.text_input(f"Titre {i+1}")
        text_input_2 = st.text_input(f"URL Concurrent {i+1}")
        dropdown = st.selectbox(f'Choisir le template de structure pour la page {i+1}', options)
        
        page_details.append({
            "page": i+1,
            "text_input_1": text_input_1,
            "text_input_2": text_input_2,
            "dropdown": dropdown
        })


    st.subheader("Cahier des charges")
    pdf_docs = st.file_uploader(
        "Upload le cahier des charges ", accept_multiple_files=False)
    st.subheader("Keywords")
    keywords = st.text_input("Enter keywords")

      
    if st.button("process"):
        with st.spinner("Processing"):
            texts = get_pdf_text_cahier(pdf_docs)

            extracted_data=extract_data("Extract alll useful data about the client information and desires "+texts )
            if not texts:
                st.write("No text extracted from PDFs.")
                return

            output_contenu=[]
            for i in range(num_pages):
                selected_template = page_details[i]['dropdown']
                with open(os.path.join(pdf_templates_folder, selected_template), 'rb') as template_file:
                    structure_text = get_pdf_text(template_file)
                output_page=handle_userinput("Provided information  : "+texts+" Write SEO content following the structure \n STRUCTURE : " + structure_text + "taking into account following keywords : "+ keywords)
                output_contenu.append(output_page)
        titles = [page["text_input_1"] for page in page_details]
        word_buffer = create_word_doc(output_contenu,titles)
        st.download_button(label="Download Word Document",
                data=word_buffer,
                file_name="SEOCONTENT.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
if __name__ == "__main__":
    main()
